from docx import Document
import os


class Parser:

    def get_data(self, file):
        doc = Document(file)
        content = self.__merge_paragraphs(doc)
        # name = self.__get_file_name(url)

        return content #, name

    def __merge_paragraphs(self, doc):
        text = ""
        counter = 0
        doc_index = 0
        space = ""

        while doc_index < len(doc.paragraphs):
            if len(doc.paragraphs[doc_index].text) == 0:
                self.__delete_paragraph(doc.paragraphs[doc_index])
                continue

            if counter > 0:
                space = " "

            text = text + space + doc.paragraphs[doc_index].text
            counter += 1
            doc_index += 1

        return text

    @staticmethod
    def __get_file_name(url):
        file_name_with_extension = os.path.basename(url)
        file_name = os.path.splitext(file_name_with_extension)[0]

        return file_name

    @staticmethod
    def __delete_paragraph(paragraph):
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None
